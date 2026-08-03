#!/usr/bin/env python3
"""
ينشئ قالب Word وزاري بصيغة docx يحتوي على placeholders متوافقة مع docxtemplater.
المفاتيح تطابق مخطط JSON للخطة حرفياً.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ضبط الهوامش
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# تعيين الخط الافتراضي
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)
# دعم RTL
rPr = style.element.get_or_add_rPr()
rtl = rPr.makeelement(qn('w:rtl'), {qn('w:val'): '1'})
rPr.append(rtl)

# ==================== العنوان ====================
title = doc.add_heading('خطة درس يومية', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x0d, 0x6b, 0x56)

# ==================== جدول المعلومات الأساسية ====================
doc.add_paragraph('')
info_heading = doc.add_heading('المعلومات الأساسية', level=1)
for run in info_heading.runs:
    run.font.color.rgb = RGBColor(0x0d, 0x6b, 0x56)

info_table = doc.add_table(rows=4, cols=4)
info_table.style = 'Table Grid'
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

info_data = [
    ("المادة", "{basic_info.subject}", "الصف", "{basic_info.grade}"),
    ("الوحدة", "{basic_info.unit}", "الدرس", "{basic_info.lesson}"),
    ("التاريخ", "{basic_info.date}", "عدد الحصص", "{basic_info.periods}"),
    ("الصفحات", "{basic_info.pages}", "", ""),
]

for i, row_data in enumerate(info_data):
    row = info_table.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(cell_text)
        run.font.size = Pt(11)
        if j % 2 == 0:  # Labels
            run.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ==================== الأهداف التعليمية ====================
doc.add_paragraph('')
obj_heading = doc.add_heading('الأهداف التعليمية', level=1)
for run in obj_heading.runs:
    run.font.color.rgb = RGBColor(0x0d, 0x6b, 0x56)

doc.add_paragraph('الأهداف المعرفية:', style='List Bullet')
# Loop placeholder for cognitive objectives
doc.add_paragraph('{#objectives.cognitive}', style='List Bullet 2')
doc.add_paragraph('الأهداف المهارية:', style='List Bullet')
doc.add_paragraph('{#objectives.skills}', style='List Bullet 2')
doc.add_paragraph('الأهداف الوجدانية:', style='List Bullet')
doc.add_paragraph('{#objectives.affective}', style='List Bullet 2')

# ==================== التهيئة ====================
doc.add_paragraph('')
warmup_heading = doc.add_heading('التهيئة', level=1)
for run in warmup_heading.runs:
    run.font.color.rgb = RGBColor(0x0d, 0x6b, 0x56)
doc.add_paragraph('{warm_up}')

# ==================== الاستراتيجيات والوسائل ====================
doc.add_paragraph('')
strat_heading = doc.add_heading('استراتيجيات التدريس والوسائل', level=1)
for run in strat_heading.runs:
    run.font.color.rgb = RGBColor(0x0d, 0x6b, 0x56)

doc.add_paragraph('الاستراتيجيات:', style='List Bullet')
doc.add_paragraph('{#strategies}', style='List Bullet 2')
doc.add_paragraph('الوسائل التعليمية:', style='List Bullet')
doc.add_paragraph('{#materials}', style='List Bullet 2')

# ==================== خطوات التنفيذ ====================
doc.add_paragraph('')
proc_heading = doc.add_heading('خطوات تنفيذ الدرس', level=1)
for run in proc_heading.runs:
    run.font.color.rgb = RGBColor(0x0d, 0x6b, 0x56)

proc_table = doc.add_table(rows=1, cols=4)
proc_table.style = 'Table Grid'
proc_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
hdr = proc_table.rows[0]
headers = ['الخطوة', 'الزمن (دقيقة)', 'دور المعلم', 'دور الطالب']
for j, h in enumerate(headers):
    cell = hdr.cells[j]
    cell.text = ""
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x0d, 0x6b, 0x56)

# Data row with loop
proc_table.add_row()
data_row = proc_table.rows[1]
data_cells = ['{#procedures}', '{#procedures.time_minutes}', '{#procedures.teacher_role}', '{#procedures.student_role}']
# Actually, docxtemplater loops need special syntax. Let's use a simpler approach.
# We'll use a single row with loop markers.
# Clear the table and rebuild with proper loop syntax
proc_table._element.getparent().remove(proc_table._element)

# Rebuild with proper loop
proc_table = doc.add_table(rows=1, cols=4)
proc_table.style = 'Table Grid'
proc_table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = proc_table.rows[0]
for j, h in enumerate(headers):
    cell = hdr.cells[j]
    cell.text = ""
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x0d, 0x6b, 0x56)

# Add loop row — docxtemplater uses {#loop} ... {/loop} syntax
row = proc_table.add_row()
row.cells[0].text = '{#procedures}{step}'
row.cells[1].text = '{time_minutes}'
row.cells[2].text = '{teacher_role}'
row.cells[3].text = '{student_role}{/procedures}'

# ==================== التقويم ====================
doc.add_paragraph('')
assess_heading = doc.add_heading('التقويم', level=1)
for run in assess_heading.runs:
    run.font.color.rgb = RGBColor(0x0d, 0x6b, 0x56)

doc.add_paragraph('التقويم القبلي:', style='List Bullet')
doc.add_paragraph('{assessment.diagnostic}')

doc.add_paragraph('التقويم البنائي:', style='List Bullet')
doc.add_paragraph('{#assessment.formative}', style='List Bullet 2')

doc.add_paragraph('التقويم الختامي:', style='List Bullet')
doc.add_paragraph('{#assessment.summative}', style='List Bullet 2')

# ==================== القيم التربوية ====================
doc.add_paragraph('')
val_heading = doc.add_heading('القيم التربوية', level=1)
for run in val_heading.runs:
    run.font.color.rgb = RGBColor(0x0d, 0x6b, 0x56)
doc.add_paragraph('{#values}', style='List Bullet')

# ==================== الدمج التكنولوجي ====================
doc.add_paragraph('')
tech_heading = doc.add_heading('الدمج التكنولوجي', level=1)
for run in tech_heading.runs:
    run.font.color.rgb = RGBColor(0x0d, 0x6b, 0x56)
doc.add_paragraph('{tech_integration}')

# ==================== مراعاة الفروق الفردية ====================
doc.add_paragraph('')
diff_heading = doc.add_heading('مراعاة الفروق الفردية', level=1)
for run in diff_heading.runs:
    run.font.color.rgb = RGBColor(0x0d, 0x6b, 0x56)

diff_table = doc.add_table(rows=2, cols=2)
diff_table.style = 'Table Grid'
diff_table.alignment = WD_TABLE_ALIGNMENT.CENTER

diff_headers = ['دعم المتعثرين', 'إثراء المتفوقين']
for j, h in enumerate(diff_headers):
    cell = diff_table.rows[0].cells[j]
    cell.text = ""
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x0d, 0x6b, 0x56)

diff_table.rows[1].cells[0].text = '{differentiation.support}'
diff_table.rows[1].cells[1].text = '{differentiation.enrichment}'

# ==================== الواجب المنزلي ====================
doc.add_paragraph('')
hw_heading = doc.add_heading('الواجب المنزلي', level=1)
for run in hw_heading.runs:
    run.font.color.rgb = RGBColor(0x0d, 0x6b, 0x56)
doc.add_paragraph('{homework}')

# ==================== الربط بالحياة ====================
doc.add_paragraph('')
life_heading = doc.add_heading('الربط بالحياة', level=1)
for run in life_heading.runs:
    run.font.color.rgb = RGBColor(0x0d, 0x6b, 0x56)
doc.add_paragraph('{real_life_connection}')

# ==================== التوقيعات ====================
doc.add_paragraph('')
doc.add_paragraph('')
sig_table = doc.add_table(rows=2, cols=3)
sig_table.style = 'Table Grid'
sig_headers = ['معلم المادة', 'رئيس القسم', 'مدير المدرسة']
for j, h in enumerate(sig_headers):
    cell = sig_table.rows[0].cells[j]
    cell.text = ""
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.size = Pt(11)
    run.font.alignment = WD_ALIGN_PARAGRAPH.CENTER
sig_table.rows[1].cells[0].text = ' '
sig_table.rows[1].cells[1].text = ' '
sig_table.rows[1].cells[2].text = ' '

# حفظ الملف
output_path = '/home/ubuntu/teacher-assistant/server/templates/plan_template_qa.docx'
import os
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"Template saved to: {output_path}")
